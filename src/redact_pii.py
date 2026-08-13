import os
import docx
from src.detector import Detector
from src.replacer import Replacer, PiiTypes


def redact_docx(input_path: str, output_path: str) -> dict:
    """
    Reads the input DOCX, redacts all detected PII, writes to output_path,
    and returns a dictionary of detection statistics.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found at: {input_path}")

    doc = docx.Document(input_path)
    detector = Detector()
    replacer = Replacer()

    # Track PII statistics
    stats_keys = [
        PiiTypes.PERSON, PiiTypes.EMAIL, PiiTypes.PHONE, PiiTypes.COMPANY,
        PiiTypes.ADDRESS, PiiTypes.SSN, PiiTypes.CREDIT_CARD, PiiTypes.DOB,
        PiiTypes.IP_ADDRESS
    ]
    counts = {key: 0 for key in stats_keys}

    def process_paragraph(p):
        if not p.text.strip():
            return
        
        matches = detector.detect(p.text)
        if not matches:
            return
        
        # Count detections
        for m in matches:
            counts[m['type']] += 1
            
        # Fallback if paragraph has text but no runs (uncommon in docx but possible)
        if not p.runs:
            new_text = p.text
            for m in sorted(matches, key=lambda x: x['start'], reverse=True):
                start, end = m['start'], m['end']
                orig = p.text[start:end]
                rep = replacer.get_replacement(orig, m['type'])
                new_text = new_text[:start] + rep + new_text[end:]
            p.text = new_text
            return

        # Build character-to-run map for precise run edits (preserves formatting)
        char_map = []
        for run in p.runs:
            for i, char in enumerate(run.text):
                char_map.append({
                    'run': run,
                    'index_in_run': i,
                    'char': char
                })
                
        # Process matches in descending order so run modifications do not offset earlier indices
        for m in sorted(matches, key=lambda x: x['start'], reverse=True):
            start = m['start']
            end = m['end']
            orig = p.text[start:end]
            rep = replacer.get_replacement(orig, m['type'])
            
            # Map character range back to the underlying runs
            runs_involved = []
            run_to_indices = {}
            for item in char_map[start:end]:
                r = item['run']
                if r not in run_to_indices:
                    runs_involved.append(r)
                    run_to_indices[r] = []
                run_to_indices[r].append(item['index_in_run'])
                
            if not runs_involved:
                continue
                
            # Replace target slice in the first run
            first_run = runs_involved[0]
            first_indices = run_to_indices[first_run]
            min_idx, max_idx = min(first_indices), max(first_indices)
            run_text_list = list(first_run.text)
            run_text_list[min_idx : max_idx + 1] = list(rep)
            first_run.text = "".join(run_text_list)
            
            # Clear target slice in any subsequent runs (e.g. if name was split across runs)
            for other_run in runs_involved[1:]:
                other_indices = run_to_indices[other_run]
                min_idx, max_idx = min(other_indices), max(other_indices)
                run_text_list = list(other_run.text)
                run_text_list[min_idx : max_idx + 1] = []
                other_run.text = "".join(run_text_list)

    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                # Cells contain paragraphs and potentially nested tables
                for p in cell.paragraphs:
                    process_paragraph(p)
                for nested_table in cell.tables:
                    process_table(nested_table)

    # 1. Process body paragraphs
    for p in doc.paragraphs:
        process_paragraph(p)

    # 2. Process body tables
    for t in doc.tables:
        process_table(t)

    # 3. Process headers and footers across all sections
    processed_sections = set()
    for section in doc.sections:
        for hf_name in ['header', 'first_page_header', 'even_page_header', 'footer', 'first_page_footer', 'even_page_footer']:
            hf = getattr(section, hf_name, None)
            if hf is not None and id(hf) not in processed_sections:
                processed_sections.add(id(hf))
                for p in hf.paragraphs:
                    process_paragraph(p)
                for t in hf.tables:
                    process_table(t)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return counts


def main():
    input_path = os.path.join("input", "source_document.docx")
    output_path = os.path.join("output", "redacted_output.docx")
    
    print("Starting PII Redaction...")
    try:
        counts = redact_docx(input_path, output_path)
        
        print("\nPII Redaction completed.\n")
        print(f"Person names detected: {counts[PiiTypes.PERSON]}")
        print(f"Emails detected: {counts[PiiTypes.EMAIL]}")
        print(f"Phone numbers detected: {counts[PiiTypes.PHONE]}")
        print(f"Companies detected: {counts[PiiTypes.COMPANY]}")
        print(f"Addresses detected: {counts[PiiTypes.ADDRESS]}")
        print(f"SSNs detected: {counts[PiiTypes.SSN]}")
        print(f"Credit cards detected: {counts[PiiTypes.CREDIT_CARD]}")
        print(f"Dates of birth detected: {counts[PiiTypes.DOB]}")
        print(f"IP addresses detected: {counts[PiiTypes.IP_ADDRESS]}")
        print(f"\nOutput:\n{output_path}")
        
    except Exception as e:
        print(f"Error during redaction: {e}")


if __name__ == "__main__":
    main()
